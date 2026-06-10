from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify
import pandas as pd
import os, io, uuid, re, sqlite3
from datetime import datetime, date
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import tempfile

app = Flask(__name__)
app.secret_key = 'invigilation_secret_2026'

UPLOAD_DIR = os.path.join(tempfile.gettempdir(), 'invig_uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

DB_PATH = 'invigilation.db'

# ── In-memory state ───────────────────────────────────────────────────────────
STATE = {
    'faculty_tt_path':    None,
    'fn_tt_path':         None,
    'an_tt_path':         None,
    'joining_dates_path': None,
    'allocations':        None,
    'counts':             {},
    'swap_requests':      [],
    'peer_swap_requests': [],
    'custom_max':         {},
    'fn_per_slot':        6,
    'an_per_slot':        6,
    'joining_dates':      {},   # {faculty_name: datetime}
    'designation_map':    {},   # {faculty_name: designation string}
}

# ── Prefix stripping ──────────────────────────────────────────────────────────
TITLE_PREFIXES = re.compile(r'^(Mrs?[.]?|Mr[.]?|Dr[.]?|Prof[.]?|Ms[.]?)\s*', re.IGNORECASE)

def strip_title(name: str) -> str:
    return TITLE_PREFIXES.sub('', name.strip()).strip()

# ── Hardcoded excluded faculty (Professors / Associate Professors) ─────────────
# These names are ALWAYS excluded from invigilation duty regardless of
# what the timetable or designation file says.
# Stored as lower-case bare names (no Mr/Mrs/Dr prefix) for robust matching.
_EXCLUDED_RAW = [
    "a. sharada",
    "m. seetha",
    "n. kalyani",
    "k. l. s. soujanya",
    "jayashree s patil",
    "d.v. lalitha parameswari",   # spelling variant 1
    "d.v. lalita parameswari",    # spelling variant 2 (single 't')
    "g. malini devi",
    "raghavender k.v.",
    "g. narendrababu reddy",      # Professor — must be excluded
]
EXCLUDED_FACULTY_NAMES = set(_EXCLUDED_RAW)

# ── Designation-based exclusion ───────────────────────────────────────────────
# Only Asst. Professors get invigilation duties.
# Professors (all grades), Associate Professors, and HoD are EXCLUDED.
def is_excluded_by_designation(faculty_name: str) -> bool:
    """Return True if the faculty should NOT be assigned invigilation duty."""

    # 1. Check hardcoded exclusion list first (strip prefix, case-insensitive)
    if strip_title(faculty_name).lower() in EXCLUDED_FACULTY_NAMES:
        return True

    # 2. Check designation map
    desig_map = STATE.get('designation_map', {})
    desig = desig_map.get(faculty_name, '')
    if not desig:
        # Try stripped-name match (handle prefix differences)
        stripped_f = strip_title(faculty_name).lower()
        for fname, fd in desig_map.items():
            if strip_title(fname).lower() == stripped_f:
                desig = fd
                break
    d_lower = desig.strip().lower()
    # Exclude: Professor, Professor & Head, Professor & Dean, Assoc. Professor
    if d_lower.startswith('professor'):
        return True
    if 'assoc' in d_lower and 'professor' in d_lower:
        return True
    if 'associate' in d_lower and 'professor' in d_lower:
        return True
    return False

# ── Column mapping ────────────────────────────────────────────────────────────
DAY_COL_MAP = {
    'MON': {'name_col': 1,  'slots': [2,3,4,5,6,7]},
    'TUE': {'name_col': 1,  'slots': [9,10,11,12,13,14]},
    'WED': {'name_col': 15, 'slots': [16,17,18,19,20,21]},
    'THU': {'name_col': 15, 'slots': [23,24,25,26,27,28]},
    'FRI': {'name_col': 29, 'slots': [30,31,32,33,34,35]},
    'SAT': {'name_col': 29, 'slots': [37,38,39,40,41,42]},
}
SLOT_LABELS  = ['9:00-10:00','10:00-11:00','11:00-12:00','1:00-2:00','2:00-3:00','3:00-4:00']
ADMIN_USER   = 'admin'
ADMIN_PASS   = 'admin123'
FACULTY_PASS = 'pass123'

# ── DOJ → max invigilation count ──────────────────────────────────────────────
def doj_to_max_count(doj: datetime) -> int:
    """
    Years of experience (from DOJ to today) → max invigilations:
      >= 25 yrs → 1
      >= 23 yrs → 2
      >= 20 yrs → 3
      >= 15 yrs → 4
      <  15 yrs → 5
    """
    today = datetime.today()
    years = (today - doj).days / 365.25
    if years >= 25:
        return 1
    elif years >= 23:
        return 2
    elif years >= 20:
        return 3
    elif years >= 15:
        return 4
    else:
        return 5

# ── Load joining dates from .docx (Word table) ────────────────────────────────
def load_joining_dates_docx(path: str) -> dict:
    """
    Reads the Staff_List Word document (.docx).
    Parses ALL tables:
      - Teaching staff table: 4 cols (S.No | Staff Details | Designation | Joining Date)
      - Non-teaching table:   3 cols (S.No | Name | Date stored in col 2)
    Returns {faculty_name: datetime} and populates STATE['designation_map'].
    """
    try:
        from docx import Document
    except ImportError:
        return {}

    try:
        doc = Document(path)
    except Exception:
        return {}

    result    = {}   # {name: datetime}
    desig_map = {}   # {name: designation}

    for table in doc.tables:
        if not table.rows:
            continue

        header_row = table.rows[0]
        headers    = [cell.text.strip().lower() for cell in header_row.cells]
        n_cols     = len(headers)

        # Detect column positions from header keywords
        name_idx  = None
        doj_idx   = None
        desig_idx = None

        for i, h in enumerate(headers):
            if name_idx is None and any(kw in h for kw in ('name', 'staff', 'faculty')):
                name_idx = i
            if doj_idx is None and any(kw in h for kw in ('join', 'doj', 'date')):
                doj_idx = i
            if desig_idx is None and any(kw in h for kw in ('desig', 'role', 'post')):
                desig_idx = i

        # Fallback column positions based on table width:
        # 4-col: S.No | Staff Details | Designation | Joining Date
        # 3-col: S.No | Name          | Date (in 3rd col, no separate desig col)
        if name_idx is None:
            name_idx = 1 if n_cols >= 2 else 0
        if doj_idx is None:
            doj_idx = 3 if n_cols >= 4 else (n_cols - 1 if n_cols >= 2 else None)
        if desig_idx is None and n_cols >= 3:
            desig_idx = 2

        if doj_idx is None:
            continue

        # Parse data rows (skip header row)
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) <= max(name_idx, doj_idx):
                continue
            name  = cells[name_idx].text.strip()
            raw   = cells[doj_idx].text.strip()
            desig = cells[desig_idx].text.strip() if (desig_idx is not None and desig_idx < len(cells)) else ''

            # Skip blank or header-repeat rows
            if not name or name.lower() in ('', 'nan'):
                continue
            if re.match(r'^s\.?\s*no\.?$', name, re.IGNORECASE):
                continue

            # Some date cells contain two dates separated by whitespace/newline
            # e.g. "19/08/2002\n13-03-2026" — always use the FIRST (original joining date)
            raw_clean  = re.split(r'[\s\n]+', raw.strip())[0]
            parsed_doj = None
            for candidate in [raw_clean, raw.strip()]:
                try:
                    parsed_doj = pd.to_datetime(candidate, dayfirst=True).to_pydatetime()
                    break
                except Exception:
                    continue

            if parsed_doj is not None:
                result[name]    = parsed_doj
                desig_map[name] = desig

    # Persist designation map to STATE so exclusion logic can use it
    STATE['designation_map'] = desig_map

    # Paragraph-based fallback for non-table formats
    if not result:
        name_pat = re.compile(r'name[:\s]+(.+)', re.IGNORECASE)
        doj_pat  = re.compile(r'(?:doj|date of joining|joining date)[:\s]+(.+)', re.IGNORECASE)
        current_name = None
        for para in doc.paragraphs:
            text = para.text.strip()
            m = name_pat.match(text)
            if m:
                current_name = m.group(1).strip()
            m2 = doj_pat.match(text)
            if m2 and current_name:
                try:
                    doj = pd.to_datetime(m2.group(1).strip(), dayfirst=True)
                    result[current_name] = doj.to_pydatetime()
                    current_name = None
                except Exception:
                    pass

    return result

# ── Load joining dates — auto-detect xlsx or docx ────────────────────────────
def load_joining_dates(path: str) -> dict:
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        return load_joining_dates_docx(path)
    # fallback: xlsx
    try:
        df = pd.read_excel(path, sheet_name=0)
    except Exception:
        return {}

    df.columns = [str(c).strip() for c in df.columns]
    col_map = {c.lower(): c for c in df.columns}

    name_key = next((col_map[k] for k in col_map if 'name' in k), None)
    doj_key  = next((col_map[k] for k in col_map
                     if 'join' in k or 'doj' in k or 'date' in k), None)

    if not name_key or not doj_key:
        cols = list(df.columns)
        if len(cols) >= 2:
            name_key = cols[0]
            doj_key  = cols[1]
        else:
            return {}

    result = {}
    for _, row in df.iterrows():
        name = str(row.get(name_key, '')).strip()
        raw  = row.get(doj_key)
        if not name or name.lower() == 'nan':
            continue
        try:
            doj = pd.to_datetime(raw)
            result[name] = doj.to_pydatetime()
        except Exception:
            continue
    return result

# ── Database ──────────────────────────────────────────────────────────────────
def init_db():
    con = sqlite3.connect(DB_PATH)
    con.execute('''CREATE TABLE IF NOT EXISTS faculty_counts (
        faculty_name    TEXT PRIMARY KEY,
        max_count       INTEGER DEFAULT 4,
        current_count   INTEGER DEFAULT 0,
        remaining_count INTEGER DEFAULT 4
    )''')
    con.commit()
    con.close()

def sync_db_from_state():
    if not STATE.get('faculty_tt_path'):
        return
    faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
    tier_max     = get_effective_max_map(faculty_list)
    counts       = STATE.get('counts') or {}
    custom_max   = STATE.get('custom_max', {})
    con = sqlite3.connect(DB_PATH)
    for f in faculty_list:
        mx        = custom_max.get(f, tier_max.get(f, 5))
        current   = counts.get(f, 0)
        remaining = max(0, mx - current)
        con.execute('''INSERT INTO faculty_counts(faculty_name,max_count,current_count,remaining_count)
                       VALUES(?,?,?,?)
                       ON CONFLICT(faculty_name) DO UPDATE SET
                         max_count=excluded.max_count,
                         current_count=excluded.current_count,
                         remaining_count=excluded.remaining_count''',
                    (f, mx, current, remaining))
    con.commit()
    con.close()

def update_faculty_in_db(faculty_name):
    if not STATE.get('faculty_tt_path'):
        return
    faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
    tier_max     = get_effective_max_map(faculty_list)
    custom_max   = STATE.get('custom_max', {})
    counts       = STATE.get('counts') or {}
    mx        = custom_max.get(faculty_name, tier_max.get(faculty_name, 5))
    current   = counts.get(faculty_name, 0)
    remaining = max(0, mx - current)
    con = sqlite3.connect(DB_PATH)
    con.execute('''INSERT INTO faculty_counts(faculty_name,max_count,current_count,remaining_count)
                   VALUES(?,?,?,?)
                   ON CONFLICT(faculty_name) DO UPDATE SET
                     max_count=excluded.max_count,
                     current_count=excluded.current_count,
                     remaining_count=excluded.remaining_count''',
                (faculty_name, mx, current, remaining))
    con.commit()
    con.close()

# ── Data helpers ──────────────────────────────────────────────────────────────
def load_faculty_list(path, eligible_only=False):
    """
    Load all faculty names from the master timetable xlsx.
    If eligible_only=True, excludes Professors / Assoc. Professors / HoD
    so only Asst. Professors, Teaching Assistants and Non-Teaching staff remain.
    NOTE: designation_map must already be populated (via load_joining_dates_docx)
    before calling with eligible_only=True for filtering to work correctly.
    """
    df = pd.read_excel(path, sheet_name=0, header=None)
    names = []
    for i in range(3, len(df)):
        n = str(df.iloc[i, 1]).strip()
        if n and n.lower() != 'nan':
            names.append(n)
    if eligible_only:
        names = [n for n in names if not is_excluded_by_designation(n)]
    return names

def load_exam_tt(path):
    df = pd.read_excel(path, sheet_name=0, header=None)
    n_rows = len(df)

    # Detect format:
    # Format A (4-row): rows 0-1 are header/label rows, row 2 = dates, row 3 = subjects
    # Format B (2-row): row 0 = dates, row 1 = subjects
    # Format C (any): scan all rows to find date row and subject row automatically
    date_row_idx    = None
    subject_row_idx = None

    # Try to auto-detect which rows contain dates and subjects
    for i in range(n_rows):
        row_vals = [str(df.iloc[i, c]).strip() for c in range(len(df.columns))]
        non_nan  = [v for v in row_vals if v and v.lower() != 'nan']
        if not non_nan:
            continue
        # Check if this row looks like a date row (try parsing first non-nan cell)
        if date_row_idx is None:
            for v in non_nan:
                # Strip day-name suffixes like "(MONDAY)" or "(TUESDAY)" etc.
                clean = re.sub(r'\s*\([^)]*\)', '', v).strip()
                try:
                    pd.to_datetime(clean, dayfirst=True)
                    date_row_idx = i
                    break
                except Exception:
                    pass
        elif subject_row_idx is None:
            # First row after date row that has non-nan content is the subject row
            subject_row_idx = i
            break

    if date_row_idx is None or subject_row_idx is None:
        return []

    exams = []
    for col in range(len(df.columns)):
        d = df.iloc[date_row_idx, col]
        s = df.iloc[subject_row_idx, col]
        if pd.notna(s) and str(s).strip() and str(s).strip().lower() != 'nan':
            # Parse date — may be a datetime object OR a string like "06/04/2026 (MONDAY)"
            try:
                if isinstance(d, str):
                    # Strip day-name in parentheses, e.g. " (MONDAY)"
                    clean_d = re.sub(r'\s*\([^)]*\)', '', d).strip()
                    dt = pd.to_datetime(clean_d, dayfirst=True)
                else:
                    dt = pd.to_datetime(d)
                exams.append({'date':    dt.strftime('%d-%m-%Y'),
                              'day':     dt.strftime('%A').upper()[:3],
                              'subject': str(s).strip()})
            except Exception:
                continue
    return exams

def get_faculty_schedule(path):
    df = pd.read_excel(path, sheet_name=0, header=None)
    schedule = {}
    for i in range(3, len(df)):
        for day, cfg in DAY_COL_MAP.items():
            name = str(df.iloc[i, cfg['name_col']]).strip()
            if not name or name.lower() == 'nan':
                continue
            schedule.setdefault(name, {})
            busy = []
            for idx, col in enumerate(cfg['slots']):
                val = str(df.iloc[i, col]).strip()
                if val and val.lower() not in ('nan', ''):
                    busy.append(SLOT_LABELS[idx])
            schedule[name][day] = busy
    return schedule

def is_free_for_exam(schedule, faculty, day, exam_type):
    busy = set(schedule.get(faculty, {}).get(day, []))
    required = ({'9:00-10:00','10:00-11:00','11:00-12:00','1:00-2:00'}
                if exam_type == 'FN' else
                {'1:00-2:00','2:00-3:00','3:00-4:00'})
    return len(required & busy) == 0

def get_effective_max_map(faculty_list):
    """
    Build {faculty_name: max_count} using DOJ if available, else fallback tier.
    """
    joining_dates = STATE.get('joining_dates', {})
    _, tier_max   = build_tier_info(faculty_list)
    result        = {}
    for f in faculty_list:
        doj = joining_dates.get(f)
        if doj is None:
            stripped_f = strip_title(f).lower()
            for jname, jdoj in joining_dates.items():
                if strip_title(jname).lower() == stripped_f:
                    doj = jdoj
                    break
        if doj is not None:
            result[f] = doj_to_max_count(doj)
        else:
            result[f] = tier_max.get(f, 5)
    return result

def build_tier_info(faculty_list):
    n  = len(faculty_list)
    t1 = n // 3
    t2 = n // 3
    tier_of = {}; tier_max = {}
    for i, f in enumerate(faculty_list):
        if i < t1:
            tier_of[f] = 1; tier_max[f] = 2
        elif i < t1 + t2:
            tier_of[f] = 2; tier_max[f] = 3
        else:
            tier_of[f] = 3; tier_max[f] = 4
    return tier_of, tier_max

def _experience_tier(eff_max_val: int) -> int:
    """
    Map a faculty member's max-count (derived from DOJ experience) to a
    priority tier.  Lower tier number = assigned FIRST (least experienced).

    max_count == 5  → experience < 15 yrs  → tier 1  (highest priority)
    max_count == 4  → experience 15–20 yrs → tier 2
    max_count == 3  → experience 20–23 yrs → tier 3
    max_count == 2  → experience 23–25 yrs → tier 4
    max_count == 1  → experience >= 25 yrs → tier 5  (lowest priority)
    """
    tier_map = {5: 1, 4: 2, 3: 3, 2: 4, 1: 5}
    return tier_map.get(eff_max_val, 1)


def generate_allocations(fn_exams, an_exams, schedule, faculty_list):
    """
    Allocate invigilators with:

    1. Strict experience-based priority (junior staff assigned first):
       Tier 1 — < 15 yrs  (max=5) → Tier 2 — 15-20 yrs (max=4) →
       Tier 3 — 20-23 yrs (max=3) → Tier 4 — 23-25 yrs (max=2) →
       Tier 5 — ≥25 yrs   (max=1)

    2. No same-day double duty: a faculty assigned to FN cannot be
       assigned to AN on the same date and vice-versa.

    3. Balanced FN/AN split: each faculty's total invigilations are
       split as evenly as possible between FN and AN sessions.
       Target: fn_target = ceil(max/2), an_target = floor(max/2)
       (or equal halves when max is even).
       A faculty is eligible for a session only if they haven't yet
       reached their per-session target — this is relaxed in fallback
       passes if needed to fill all slots.
    """
    eff_max    = get_effective_max_map(faculty_list)
    custom_max = STATE.get('custom_max', {})
    for f in faculty_list:
        if f in custom_max:
            eff_max[f] = custom_max[f]

    fn_slots = STATE.get('fn_per_slot', 6)
    an_slots = STATE.get('an_per_slot', 6)

    assigned_count    = {f: 0 for f in faculty_list}
    fn_count          = {f: 0 for f in faculty_list}   # FN assignments per faculty
    an_count          = {f: 0 for f in faculty_list}   # AN assignments per faculty
    allocations       = []
    date_assigned     = {f: set() for f in faculty_list}  # dates already assigned

    # Per-faculty session targets for balanced FN/AN split
    import math
    fn_target = {f: math.ceil(eff_max[f] / 2) for f in faculty_list}
    an_target = {f: math.floor(eff_max[f] / 2) for f in faculty_list}

    # Pre-compute experience tier (static for whole run)
    faculty_tier = {f: _experience_tier(eff_max.get(f, 5)) for f in faculty_list}

    def priority_key(f):
        """Lower tier (junior) first → then fewest assignments → alpha tie-break."""
        return (faculty_tier[f], assigned_count[f], f)

    # Sort exams: process each date's FN then AN together so same-day
    # exclusion is applied correctly as we go.
    all_exams = [('FN', e, fn_slots) for e in fn_exams] + \
                [('AN', e, an_slots) for e in an_exams]
    all_exams.sort(key=lambda x: (x[1]['date'], x[0]))  # 'AN' < 'FN' alphabetically, so FN first after sort flip
    # Ensure FN is processed before AN on the same date so same-day block works
    all_exams.sort(key=lambda x: (x[1]['date'], 0 if x[0] == 'FN' else 1))

    for session_type, exam, n_needed in all_exams:
        exam_date  = exam['date']
        exam_day   = exam['day']
        chosen     = []

        def session_target_ok(f, stype):
            """True if faculty hasn't yet reached their target for this session type."""
            if stype == 'FN':
                return fn_count[f] < fn_target[f]
            else:
                return an_count[f] < an_target[f]

        # ── Pass 1 (ideal): free timetable slot, under total max,
        #    not assigned this date, within session target
        p1 = [f for f in faculty_list
              if is_free_for_exam(schedule, f, exam_day, session_type)
              and assigned_count[f] < eff_max[f]
              and exam_date not in date_assigned[f]
              and session_target_ok(f, session_type)]
        p1.sort(key=priority_key)
        chosen = p1[:n_needed]

        # ── Pass 2: relax session-balance target (still free, not same-date)
        if len(chosen) < n_needed:
            already = set(chosen)
            p2 = [f for f in faculty_list
                  if f not in already
                  and is_free_for_exam(schedule, f, exam_day, session_type)
                  and assigned_count[f] < eff_max[f]
                  and exam_date not in date_assigned[f]]
            p2.sort(key=priority_key)
            chosen += p2[:n_needed - len(chosen)]

        # ── Pass 3: relax timetable-free constraint, same-day rule ALWAYS enforced
        if len(chosen) < n_needed:
            already = set(chosen)
            p3 = [f for f in faculty_list
                  if f not in already
                  and assigned_count[f] < eff_max[f]
                  and exam_date not in date_assigned[f]]
            p3.sort(key=priority_key)
            chosen += p3[:n_needed - len(chosen)]

        # ── Pass 4: no-op — same-day double duty is NEVER allowed under any circumstances

        for f in chosen:
            allocations.append({
                'date':    exam['date'],
                'day':     exam['day'],
                'session': session_type,
                'faculty': f,
                'subject': exam['subject'],
            })
            assigned_count[f] += 1
            date_assigned[f].add(exam_date)
            if session_type == 'FN':
                fn_count[f] += 1
            else:
                an_count[f] += 1

    return allocations, assigned_count

def find_free_faculty_for_slot(date, day, session_type, exclude_faculty=None):
    """
    Find faculty who are free and eligible to take an invigilation in the given
    date/session slot.  Used for peer-swap dropdowns and admin-swap candidates.

    Key fix: we only exclude faculty already assigned to THIS EXACT SLOT.
    We do NOT exclude faculty assigned to the other session on the same date,
    because in a swap the requester is giving up their slot — the replacement
    is not being assigned a new duty on top of an existing same-day duty;
    they are simply filling the requester's vacated slot.

    Faculty under their max count who are free during the exam session window
    and not already in this slot are shown.
    """
    if not STATE.get('faculty_tt_path'):
        return []
    allocs       = STATE.get('allocations') or []
    faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
    schedule     = get_faculty_schedule(STATE['faculty_tt_path'])
    counts       = STATE.get('counts') or {}
    eff_max      = get_effective_max_map(faculty_list)
    custom_max   = STATE.get('custom_max', {})
    for f in faculty_list:
        if f in custom_max:
            eff_max[f] = custom_max[f]

    # Only exclude faculty already IN this exact slot (date + session match).
    # Do NOT exclude faculty assigned to the other session on the same date —
    # they are free for this slot from a scheduling perspective.
    already_this_slot = {a['faculty'] for a in allocs
                         if a['date'] == date and a['session'] == session_type}

    candidates = []
    for f in faculty_list:
        if exclude_faculty and f == exclude_faculty:
            continue
        if f in already_this_slot:
            continue
        if counts.get(f, 0) >= eff_max.get(f, 5):
            continue
        if is_free_for_exam(schedule, f, day, session_type):
            candidates.append(f)

    # Also include faculty who are timetable-busy but still under max and not in slot
    # (relaxed fallback so the dropdown is never empty when capacity exists)
    if not candidates:
        for f in faculty_list:
            if exclude_faculty and f == exclude_faculty:
                continue
            if f in already_this_slot:
                continue
            if counts.get(f, 0) >= eff_max.get(f, 5):
                continue
            if f not in candidates:
                candidates.append(f)

    candidates.sort(key=lambda f: (_experience_tier(eff_max.get(f, 5)), counts.get(f, 0), f))
    return candidates

# ── Routes ────────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    msg = None
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        role     = request.form.get('role', '').strip()

        if role == 'admin':
            if username == ADMIN_USER and password == ADMIN_PASS:
                session['user'] = username
                session['role'] = 'admin'
                return redirect(url_for('admin_dashboard'))
            msg = 'Invalid admin credentials.'

        elif role == 'faculty':
            if password != FACULTY_PASS:
                msg = 'Incorrect password.'
            elif not STATE['faculty_tt_path']:
                msg = 'Faculty data not loaded. Ask admin to upload the timetable first.'
            else:
                try:
                    # Use full list (not eligible_only) so ALL faculty can log in,
                    # including Professors and HoD who just won't get invig duties.
                    faculty_list = load_faculty_list(STATE['faculty_tt_path'])
                except Exception as e:
                    msg = f'Error reading faculty data: {e}'
                    return render_template('login.html', msg=msg)

                stripped_input = strip_title(username).lower()
                matched = None
                for f in faculty_list:
                    # Match without Mr/Mrs/Dr prefix
                    if strip_title(f).lower() == stripped_input:
                        matched = f
                        break
                    # Also allow full exact match with prefix
                    if f.strip().lower() == username.lower():
                        matched = f
                        break

                if matched:
                    session['user'] = matched
                    session['role'] = 'faculty'
                    return redirect(url_for('faculty_dashboard'))
                else:
                    msg = f'Name "{username}" not found in faculty list.'
        else:
            msg = 'Please select a role.'
    return render_template('login.html', msg=msg)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# ── Admin routes ──────────────────────────────────────────────────────────────
@app.route('/admin')
def admin_dashboard():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    pending_swaps = [s for s in STATE['swap_requests'] if s['status'] == 'pending']
    pending_peer  = [s for s in STATE.get('peer_swap_requests', []) if s['status'] == 'pending']
    return render_template('admin.html',
        faculty_tt_uploaded     = bool(STATE['faculty_tt_path']),
        fn_tt_uploaded          = bool(STATE['fn_tt_path']),
        an_tt_uploaded          = bool(STATE['an_tt_path']),
        joining_dates_uploaded  = bool(STATE['joining_dates_path']),
        allocations             = STATE['allocations'],
        pending_swaps           = pending_swaps,
        pending_peer            = pending_peer,
        fn_per_slot             = STATE.get('fn_per_slot', 6),
        an_per_slot             = STATE.get('an_per_slot', 6))

@app.route('/upload', methods=['POST'])
def upload():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    uploaded, errors = [], []

    def save_file(fobj, key, label, allowed_exts):
        if fobj and fobj.filename:
            ext = os.path.splitext(fobj.filename)[1].lower()
            if ext not in allowed_exts:
                errors.append(f'{label}: only {"/".join(allowed_exts)} files accepted.')
                return
            path = os.path.join(UPLOAD_DIR, f'invig_{key}{ext}')
            fobj.save(path)
            STATE[f'{key}_path'] = path
            uploaded.append(label)

    save_file(request.files.get('faculty_tt'),    'faculty_tt',    'Faculty Timetable',   ['.xlsx'])
    save_file(request.files.get('fn_tt'),         'fn_tt',         'FN Exam Timetable',   ['.xlsx'])
    save_file(request.files.get('an_tt'),         'an_tt',         'AN Exam Timetable',   ['.xlsx'])
    save_file(request.files.get('joining_dates'), 'joining_dates', 'Staff Joining Dates', ['.docx', '.xlsx'])

    # Parse joining dates immediately after upload — silently, no confusing errors
    if STATE.get('joining_dates_path'):
        jd = load_joining_dates(STATE['joining_dates_path'])
        STATE['joining_dates'] = jd
        # Only show success/info if the file was freshly uploaded in this request
        if 'Staff Joining Dates' in uploaded and jd:
            flash(f'Joining dates loaded for {len(jd)} staff members.', 'success')
        # Do NOT show an error if jd is empty — file is saved and allocations
        # will simply fall back to tier-based limits when generated.

    for e in errors:
        flash(e, 'error')
    if uploaded:
        flash('Uploaded: ' + ', '.join(uploaded), 'success')
    elif not errors:
        flash('No files selected.', 'error')
    return redirect(url_for('admin_dashboard'))

@app.route('/generate', methods=['POST'])
def generate():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    if not all([STATE['faculty_tt_path'], STATE['fn_tt_path'], STATE['an_tt_path']]):
        flash('Upload all three required files before generating.', 'error')
        return redirect(url_for('admin_dashboard'))

    try:
        fn_s = int(request.form.get('fn_per_slot', STATE.get('fn_per_slot', 6)))
        an_s = int(request.form.get('an_per_slot', STATE.get('an_per_slot', 6)))
        STATE['fn_per_slot'] = max(1, fn_s)
        STATE['an_per_slot'] = max(1, an_s)
    except ValueError:
        pass

    try:
        fn_exams     = load_exam_tt(STATE['fn_tt_path'])
        an_exams     = load_exam_tt(STATE['an_tt_path'])
        schedule     = get_faculty_schedule(STATE['faculty_tt_path'])
        # eligible_only=True: only Asst. Professors / Teaching Assistants / Non-Teaching staff
        faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
        allocs, counts = generate_allocations(fn_exams, an_exams, schedule, faculty_list)
        STATE['allocations']        = allocs
        STATE['counts']             = counts
        STATE['swap_requests']      = []
        STATE['peer_swap_requests'] = []
        init_db()
        sync_db_from_state()
        flash(f'Allocations generated! {len(allocs)} assignments across {len(faculty_list)} eligible faculty.', 'success')
    except Exception as e:
        flash(f'Error: {e}', 'error')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/exam_dates_json')
def exam_dates_json():
    """Return all valid exam dates from FN and AN timetables as YYYY-MM-DD strings."""
    if session.get('role') != 'admin':
        return jsonify({'dates': []})
    dates = set()
    for key in ('fn_tt_path', 'an_tt_path'):
        path = STATE.get(key)
        if path:
            try:
                exams = load_exam_tt(path)
                for e in exams:
                    dt = datetime.strptime(e['date'], '%d-%m-%Y')
                    dates.add(dt.strftime('%Y-%m-%d'))
            except Exception:
                pass
    return jsonify({'dates': sorted(dates)})

@app.route('/admin/slot_info')
def slot_info():
    """Return how many invigilators are already assigned for a given date+session,
    and how many free faculty (under max, not in slot, not same-day) are available."""
    if session.get('role') != 'admin':
        return jsonify({'error': 'unauthorized'}), 403
    raw_date  = request.args.get('date', '').strip()   # YYYY-MM-DD
    session_t = request.args.get('session', '').strip().upper()

    if not raw_date or session_t not in ('FN', 'AN'):
        return jsonify({'current': 0, 'available': 0})

    try:
        dt = datetime.strptime(raw_date, '%Y-%m-%d')
    except ValueError:
        return jsonify({'current': 0, 'available': 0})

    exam_date = dt.strftime('%d-%m-%Y')
    day       = dt.strftime('%A').upper()[:3]
    allocs    = STATE.get('allocations') or []
    counts    = STATE.get('counts') or {}

    already_this_slot = {a['faculty'] for a in allocs
                         if a['date'] == exam_date and a['session'] == session_t}
    already_this_date = {a['faculty'] for a in allocs if a['date'] == exam_date}
    current_in_slot   = len(already_this_slot)

    free_count = 0
    if STATE.get('faculty_tt_path'):
        faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
        eff_max      = get_effective_max_map(faculty_list)
        custom_max   = STATE.get('custom_max', {})
        for f in faculty_list:
            if f in custom_max:
                eff_max[f] = custom_max[f]
        # Count faculty who could still be added to this slot
        free_count = sum(
            1 for f in faculty_list
            if f not in already_this_slot
            and f not in already_this_date
            and counts.get(f, 0) < eff_max.get(f, 5)
        )

    return jsonify({'current': current_in_slot, 'available': free_count})

# ── Add invigilations for a specific date/session ─────────────────────────────
@app.route('/admin/add_invigilation', methods=['GET'])
def add_invigilation_page():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    allocs       = STATE.get('allocations') or []
    faculty_list = []
    if STATE.get('faculty_tt_path'):
        faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
    existing_dates = sorted(set(a['date'] for a in allocs))
    today = datetime.today().strftime('%Y-%m-%d')

    # Collect valid exam dates from both timetables (YYYY-MM-DD for input[type=date])
    valid_exam_dates = []
    for key in ('fn_tt_path', 'an_tt_path'):
        path = STATE.get(key)
        if path:
            try:
                exams = load_exam_tt(path)
                for e in exams:
                    dt  = datetime.strptime(e['date'], '%d-%m-%Y')
                    iso = dt.strftime('%Y-%m-%d')
                    if iso not in valid_exam_dates:
                        valid_exam_dates.append(iso)
            except Exception:
                pass
    valid_exam_dates = sorted(set(valid_exam_dates))

    return render_template('add_invigilations.html',
        faculty_list     = faculty_list,
        existing_dates   = existing_dates,
        valid_exam_dates = valid_exam_dates,
        fn_per_slot      = STATE.get('fn_per_slot', 6),
        an_per_slot      = STATE.get('an_per_slot', 6),
        today            = today)

@app.route('/admin/add_invigilation', methods=['POST'])
def add_invigilation_submit():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))

    raw_date   = request.form.get('exam_date', '').strip()
    session_t  = request.form.get('session', '').strip().upper()
    n_invig    = request.form.get('n_invig', '').strip()

    errors = []
    if not raw_date:
        errors.append('Please select a date.')
    if session_t not in ('FN', 'AN'):
        errors.append('Please select a valid session (FN or AN).')
    try:
        n_invig = int(n_invig)
        if n_invig < 1:
            raise ValueError
    except (ValueError, TypeError):
        errors.append('Number of invigilators must be a positive integer.')

    if errors:
        for e in errors:
            flash(e, 'error')
        return redirect(url_for('add_invigilation_page'))

    try:
        dt = datetime.strptime(raw_date, '%Y-%m-%d')
    except ValueError:
        flash('Invalid date format.', 'error')
        return redirect(url_for('add_invigilation_page'))

    exam_date = dt.strftime('%d-%m-%Y')
    day       = dt.strftime('%A').upper()[:3]

    if not STATE.get('faculty_tt_path'):
        flash('Faculty timetable not uploaded. Cannot find eligible faculty.', 'error')
        return redirect(url_for('add_invigilation_page'))

    faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
    schedule     = get_faculty_schedule(STATE['faculty_tt_path'])
    eff_max      = get_effective_max_map(faculty_list)
    custom_max   = STATE.get('custom_max', {})
    for f in faculty_list:
        if f in custom_max:
            eff_max[f] = custom_max[f]

    counts    = STATE.get('counts') or {}
    allocs    = STATE.get('allocations') or []

    # Who is already assigned to this exact slot?
    already_this_slot = {a['faculty'] for a in allocs
                         if a['date'] == exam_date and a['session'] == session_t}
    # Who is already assigned anywhere on this date (FN or AN)?
    already_this_date = {a['faculty'] for a in allocs if a['date'] == exam_date}

    # How many are already assigned in this slot?
    current_in_slot = len(already_this_slot)

    # The requested total for this slot is n_invig.
    # We only need to add the difference beyond what's already assigned.
    need_to_add = n_invig - current_in_slot

    if need_to_add <= 0:
        flash(
            f'✅ {current_in_slot} invigilator(s) are already assigned to '
            f'{exam_date} {session_t}. No additional invigilators needed.',
            'success'
        )
        return redirect(url_for('add_invigilation_page'))

    # Build eligible list: free timetable slot, under max, not already in slot,
    # not assigned to the other session on the same day (no same-day double duty)
    eligible = [f for f in faculty_list
                if is_free_for_exam(schedule, f, day, session_t)
                and counts.get(f, 0) < eff_max.get(f, 5)
                and f not in already_this_slot
                and f not in already_this_date]
    eligible.sort(key=lambda f: (_experience_tier(eff_max.get(f, 5)), counts.get(f, 0), f))
    chosen = eligible[:need_to_add]

    # Fallback: relax timetable-free constraint but NEVER allow same-day double duty
    if len(chosen) < need_to_add:
        already_chosen = set(chosen)
        fallback = [f for f in faculty_list
                    if f not in already_chosen
                    and f not in already_this_slot
                    and f not in already_this_date   # same-day rule ALWAYS enforced
                    and counts.get(f, 0) < eff_max.get(f, 5)]
        fallback.sort(key=lambda f: (_experience_tier(eff_max.get(f, 5)), counts.get(f, 0), f))
        chosen += fallback[:need_to_add - len(chosen)]

    if not chosen:
        # Distinguish: were there faculty in the pool at all, or is the pool truly exhausted?
        pool_all = [f for f in faculty_list
                    if f not in already_this_slot
                    and counts.get(f, 0) < eff_max.get(f, 5)]
        if not pool_all:
            flash(
                f'⚠️ All free faculty in this slot are already allotted and there are '
                f'no more free faculty at this slot ({exam_date} {session_t}).',
                'error'
            )
        else:
            flash(
                f'⚠️ All free faculty in this slot are already allotted and there are '
                f'no more free faculty at this slot ({exam_date} {session_t}).',
                'error'
            )
        return redirect(url_for('add_invigilation_page'))

    if STATE['allocations'] is None:
        STATE['allocations'] = []

    for f in chosen:
        STATE['allocations'].append({
            'date':    exam_date,
            'day':     day,
            'session': session_t,
            'faculty': f,
            'subject': 'Exam',
        })
        counts[f] = counts.get(f, 0) + 1
        update_faculty_in_db(f)

    STATE['counts'] = counts

    added     = len(chosen)
    total_now = current_in_slot + added
    shortage  = need_to_add - added

    if shortage > 0:
        flash(
            f'⚠️ Added {added} invigilator(s) for {exam_date} {session_t} '
            f'(total now: {total_now}). '
            f'{shortage} more requested but no additional free faculty available — '
            f'all eligible faculty at this slot are already allotted.',
            'error'
        )
    else:
        flash(
            f'✅ Added {added} invigilator(s) for {exam_date} {session_t} '
            f'(total in slot: {total_now}).',
            'success'
        )
    return redirect(url_for('add_invigilation_page'))

@app.route('/admin/modify')
def admin_modify():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    allocs = STATE.get('allocations') or []
    if not allocs:
        flash('Generate allocations first.', 'error')
        return redirect(url_for('admin_dashboard'))
    counts           = STATE.get('counts') or {}
    assigned_faculty = sorted(set(a['faculty'] for a in allocs))
    faculty_list     = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True) if STATE['faculty_tt_path'] else []
    eff_max          = get_effective_max_map(faculty_list)
    custom_max       = STATE.get('custom_max', {})
    faculty_data     = []
    for f in assigned_faculty:
        faculty_data.append({
            'name':          f,
            'current_count': counts.get(f, 0),
            'default_max':   eff_max.get(f, 5),
            'custom_max':    custom_max.get(f, ''),
        })
    return render_template('admin_modify.html', faculty_data=faculty_data)

@app.route('/admin/modify_submit', methods=['POST'])
def admin_modify_submit():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    selected   = request.form.getlist('selected_faculty')
    custom_max = {}
    for fname in selected:
        max_val = request.form.get(f'max_{fname}', '').strip()
        if max_val.isdigit() and int(max_val) >= 0:
            custom_max[fname] = int(max_val)
    STATE['custom_max'] = custom_max

    if not all([STATE['faculty_tt_path'], STATE['fn_tt_path'], STATE['an_tt_path']]):
        flash('File paths missing, cannot regenerate.', 'error')
        return redirect(url_for('admin_dashboard'))
    try:
        fn_exams     = load_exam_tt(STATE['fn_tt_path'])
        an_exams     = load_exam_tt(STATE['an_tt_path'])
        schedule     = get_faculty_schedule(STATE['faculty_tt_path'])
        faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
        allocs, counts = generate_allocations(fn_exams, an_exams, schedule, faculty_list)
        STATE['allocations']        = allocs
        STATE['counts']             = counts
        STATE['swap_requests']      = []
        STATE['peer_swap_requests'] = []
        sync_db_from_state()
        flash(f'Allocations updated! {len(allocs)} assignments generated.', 'success')
    except Exception as e:
        flash(f'Error regenerating: {e}', 'error')
    return redirect(url_for('admin_dashboard'))

@app.route('/assigned_count')
def assigned_count():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    faculty_list = []
    eff_max      = {}
    if STATE.get('faculty_tt_path'):
        faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True)
        eff_max      = get_effective_max_map(faculty_list)
    custom_max = STATE.get('custom_max', {})
    for f in faculty_list:
        if f in custom_max:
            eff_max[f] = custom_max[f]
    return render_template('assigned_count.html',
        counts   = STATE.get('counts', {}),
        eff_max  = eff_max,
        total    = len(STATE.get('allocations') or []))

@app.route('/finalize')
def finalize():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    allocs = STATE.get('allocations') or []
    if not allocs:
        flash('Generate allocations first.', 'error')
        return redirect(url_for('admin_dashboard'))
    pivot = {}
    for a in allocs:
        key = (a['date'], a['day'])
        pivot.setdefault(key, {'FN': [], 'AN': []})
        pivot[key][a['session']].append(a['faculty'])
    rows = [{'date': k[0], 'day': k[1], 'fn': pivot[k]['FN'], 'an': pivot[k]['AN']}
            for k in sorted(pivot, key=lambda k: datetime.strptime(k[0], '%d-%m-%Y'))]
    return render_template('finalize.html', rows=rows)

@app.route('/download_report')
def download_report():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    allocs = STATE.get('allocations') or []
    counts = STATE.get('counts') or {}

    wb       = Workbook()
    hdr_fill = PatternFill("solid", fgColor="1A1A6E")
    hdr_font = Font(bold=True, color="FFFFFF", size=11)
    thin     = Side(style='thin')
    bdr      = Border(left=thin, right=thin, top=thin, bottom=thin)
    ctr      = Alignment(horizontal='center', vertical='center', wrap_text=True)

    ws1 = wb.active
    ws1.title = "Invigilation Schedule"
    ws1.row_dimensions[1].height = 32
    for c, h in enumerate(['Date', 'Day', 'FN Faculty (10:00 AM - 12:00 PM)', 'AN Faculty (2:00 PM - 4:00 PM)'], 1):
        cell = ws1.cell(row=1, column=c, value=h)
        cell.font = hdr_font; cell.fill = hdr_fill
        cell.alignment = ctr; cell.border = bdr

    pivot = {}
    for a in allocs:
        key = (a['date'], a['day'])
        pivot.setdefault(key, {'FN': [], 'AN': []})
        pivot[key][a['session']].append(a['faculty'])

    fn_fill = PatternFill("solid", fgColor="DDEEFF")
    an_fill = PatternFill("solid", fgColor="FFE8D6")
    sk = sorted(pivot, key=lambda k: datetime.strptime(k[0], '%d-%m-%Y'))
    for r, k in enumerate(sk, 2):
        fn_names = '\n'.join(pivot[k]['FN']) if pivot[k]['FN'] else '-'
        an_names = '\n'.join(pivot[k]['AN']) if pivot[k]['AN'] else '-'
        row_h = max(15, 15 * max(len(pivot[k]['FN']), len(pivot[k]['AN']), 1))
        ws1.row_dimensions[r].height = row_h
        for c, val in enumerate([k[0], k[1], fn_names, an_names], 1):
            cell = ws1.cell(row=r, column=c, value=val)
            cell.border = bdr; cell.alignment = ctr
            if c == 3 and pivot[k]['FN']:
                cell.fill = fn_fill
            elif c == 4 and pivot[k]['AN']:
                cell.fill = an_fill
    ws1.column_dimensions['A'].width = 16
    ws1.column_dimensions['B'].width = 10
    ws1.column_dimensions['C'].width = 42
    ws1.column_dimensions['D'].width = 42

    ws2 = wb.create_sheet("Faculty Assignment Count")
    for c, h in enumerate(['Faculty Name', 'Max Count (DOJ-based)', 'Current Invigilations', 'Remaining'], 1):
        cell = ws2.cell(row=1, column=c, value=h)
        cell.font = hdr_font; cell.fill = hdr_fill
        cell.alignment = ctr; cell.border = bdr
    faculty_list = load_faculty_list(STATE['faculty_tt_path'], eligible_only=True) if STATE.get('faculty_tt_path') else []
    eff_max      = get_effective_max_map(faculty_list)
    custom_max   = STATE.get('custom_max', {})
    for f in faculty_list:
        if f in custom_max:
            eff_max[f] = custom_max[f]
    for r, (f, cnt) in enumerate(sorted(counts.items(), key=lambda x: -x[1]), 2):
        mx = eff_max.get(f, 5)
        for c, val in enumerate([f, mx, cnt, max(0, mx - cnt)], 1):
            cell = ws2.cell(row=r, column=c, value=val)
            cell.border = bdr; cell.alignment = ctr
    ws2.column_dimensions['A'].width = 38
    ws2.column_dimensions['B'].width = 22
    ws2.column_dimensions['C'].width = 24
    ws2.column_dimensions['D'].width = 14

    joining_dates = STATE.get('joining_dates', {})
    if joining_dates:
        ws3 = wb.create_sheet("DOJ & Max Count")
        for c, h in enumerate(['Faculty Name', 'Date of Joining', 'Years of Experience', 'Max Invigilations'], 1):
            cell = ws3.cell(row=1, column=c, value=h)
            cell.font = hdr_font; cell.fill = hdr_fill
            cell.alignment = ctr; cell.border = bdr
        today = datetime.today()
        for r, (fname, doj) in enumerate(sorted(joining_dates.items()), 2):
            yrs = round((today - doj).days / 365.25, 1)
            mx  = doj_to_max_count(doj)
            for c, val in enumerate([fname, doj.strftime('%d-%m-%Y'), yrs, mx], 1):
                cell = ws3.cell(row=r, column=c, value=val)
                cell.border = bdr; cell.alignment = ctr
        ws3.column_dimensions['A'].width = 38
        ws3.column_dimensions['B'].width = 18
        ws3.column_dimensions['C'].width = 22
        ws3.column_dimensions['D'].width = 18

    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    return send_file(buf, as_attachment=True, download_name='Invigilation_Report.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

# ── Admin swap management ─────────────────────────────────────────────────────
@app.route('/admin/swaps')
def admin_swaps():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    return render_template('admin_swaps.html',
        swap_requests      = STATE['swap_requests'],
        peer_swap_requests = STATE.get('peer_swap_requests', []))

@app.route('/admin/swap_candidates/<swap_id>')
def admin_swap_candidates(swap_id):
    if session.get('role') != 'admin':
        return jsonify({'error': 'unauthorized'}), 403
    swap = next((s for s in STATE['swap_requests'] if s['id'] == swap_id), None)
    if not swap:
        return jsonify({'error': 'not found'}), 404
    candidates = find_free_faculty_for_slot(
        swap['date'], swap['day'], swap['session'], exclude_faculty=swap['requester'])
    return jsonify({'candidates': candidates})

@app.route('/admin/swap_action', methods=['POST'])
def admin_swap_action():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))

    swap_id    = request.form.get('swap_id')
    action     = request.form.get('action')
    chosen_fac = request.form.get('chosen_faculty', '').strip()

    swap = next((s for s in STATE['swap_requests'] if s['id'] == swap_id), None)
    if not swap:
        flash('Swap request not found.', 'error')
        return redirect(url_for('admin_swaps'))

    if action == 'reject':
        swap['status'] = 'rejected'
        flash(f'Swap request from {swap["requester"]} rejected.', 'success')
        return redirect(url_for('admin_swaps'))

    if action == 'replace':
        if not chosen_fac:
            flash('Please select a replacement faculty.', 'error')
            return redirect(url_for('admin_swaps'))

        allocs    = STATE.get('allocations') or []
        counts    = STATE.get('counts') or {}
        alloc_idx = swap['alloc_index']

        if alloc_idx >= len(allocs):
            flash('Original allocation no longer valid.', 'error')
            swap['status'] = 'rejected'
            return redirect(url_for('admin_swaps'))

        alloc = allocs[alloc_idx]
        if alloc['faculty'] != swap['requester']:
            flash('Allocation has changed since request was made.', 'error')
            swap['status'] = 'rejected'
            return redirect(url_for('admin_swaps'))

        old_faculty = alloc['faculty']
        allocs[alloc_idx]['faculty'] = chosen_fac
        counts[old_faculty] = max(0, counts.get(old_faculty, 1) - 1)
        counts[chosen_fac]  = counts.get(chosen_fac, 0) + 1
        swap['status']      = 'approved'
        swap['replaced_by'] = chosen_fac
        update_faculty_in_db(old_faculty)
        update_faculty_in_db(chosen_fac)
        flash(f'Done: {old_faculty} replaced by {chosen_fac} for {alloc["date"]} {alloc["session"]}.', 'success')

    return redirect(url_for('admin_swaps'))

# ── Peer swap routes ──────────────────────────────────────────────────────────
@app.route('/faculty/peer_swap_request', methods=['POST'])
def peer_swap_request():
    if session.get('role') != 'faculty':
        return redirect(url_for('login'))
    requester   = session.get('user')
    alloc_index = int(request.form.get('alloc_index'))
    target_fac  = request.form.get('target_faculty', '').strip()

    allocs = STATE.get('allocations') or []
    if alloc_index >= len(allocs) or allocs[alloc_index]['faculty'] != requester:
        flash('Invalid allocation selected.', 'error')
        return redirect(url_for('faculty_dashboard'))
    if not target_fac:
        flash('Please select a faculty to request swap with.', 'error')
        return redirect(url_for('faculty_dashboard'))

    existing = next((s for s in STATE.get('peer_swap_requests', [])
                     if s['alloc_index'] == alloc_index and s['status'] == 'pending'), None)
    if existing:
        flash('A peer swap request for this duty already exists and is pending.', 'error')
        return redirect(url_for('faculty_dashboard'))

    alloc = allocs[alloc_index]
    STATE.setdefault('peer_swap_requests', []).append({
        'id':          str(uuid.uuid4())[:8],
        'requester':   requester,
        'target':      target_fac,
        'alloc_index': alloc_index,
        'date':        alloc['date'],
        'day':         alloc['day'],
        'session':     alloc['session'],
        'status':      'pending',
        'submitted':   datetime.now().strftime('%d-%m-%Y %H:%M'),
        'resolved':    None,
    })
    flash(f'Swap request sent to {target_fac}. Waiting for their acceptance.', 'success')
    return redirect(url_for('faculty_dashboard'))

@app.route('/faculty/peer_swap_respond', methods=['POST'])
def peer_swap_respond():
    if session.get('role') != 'faculty':
        return redirect(url_for('login'))
    faculty_name = session.get('user')
    swap_id      = request.form.get('swap_id')
    action       = request.form.get('action')

    swap = next((s for s in STATE.get('peer_swap_requests', []) if s['id'] == swap_id), None)
    if not swap or swap['target'] != faculty_name:
        flash('Invalid swap request.', 'error')
        return redirect(url_for('faculty_dashboard'))

    if action == 'decline':
        swap['status']   = 'rejected'
        swap['resolved'] = datetime.now().strftime('%d-%m-%Y %H:%M')
        flash('Swap request declined.', 'success')
        return redirect(url_for('faculty_dashboard'))

    if action == 'accept':
        allocs    = STATE.get('allocations') or []
        counts    = STATE.get('counts') or {}
        alloc_idx = swap['alloc_index']

        if alloc_idx >= len(allocs) or allocs[alloc_idx]['faculty'] != swap['requester']:
            flash('Original allocation has changed. Cannot complete swap.', 'error')
            swap['status'] = 'rejected'
            return redirect(url_for('faculty_dashboard'))

        old_faculty = swap['requester']
        new_faculty = faculty_name
        allocs[alloc_idx]['faculty'] = new_faculty
        counts[old_faculty] = max(0, counts.get(old_faculty, 1) - 1)
        counts[new_faculty] = counts.get(new_faculty, 0) + 1
        swap['status']   = 'accepted'
        swap['resolved'] = datetime.now().strftime('%d-%m-%Y %H:%M')
        update_faculty_in_db(old_faculty)
        update_faculty_in_db(new_faculty)

        STATE['swap_requests'].append({
            'id':          str(uuid.uuid4())[:8],
            'requester':   old_faculty,
            'alloc_index': alloc_idx,
            'date':        swap['date'],
            'day':         swap['day'],
            'session':     swap['session'],
            'reason':      f'Peer swap - {old_faculty} requested {new_faculty}',
            'status':      'approved',
            'replaced_by': new_faculty,
            'submitted':   swap['submitted'],
            'type':        'peer',
        })
        flash(f'Swap accepted! You are now assigned to {swap["date"]} {swap["session"]} instead of {old_faculty}.', 'success')

    return redirect(url_for('faculty_dashboard'))

# ── Faculty routes ────────────────────────────────────────────────────────────
@app.route('/faculty')
def faculty_dashboard():
    if session.get('role') != 'faculty':
        return redirect(url_for('login'))
    faculty_name = session.get('user')
    allocs    = STATE.get('allocations') or []
    my_allocs = [(i, a) for i, a in enumerate(allocs) if a['faculty'] == faculty_name]
    my_swaps  = [s for s in STATE['swap_requests'] if s['requester'] == faculty_name]

    peer_sent     = [s for s in STATE.get('peer_swap_requests', []) if s['requester'] == faculty_name]
    peer_incoming = [s for s in STATE.get('peer_swap_requests', [])
                     if s['target'] == faculty_name and s['status'] == 'pending']

    eligible_map = {}
    if STATE.get('faculty_tt_path'):
        for idx, a in my_allocs:
            eligible = find_free_faculty_for_slot(
                a['date'], a['day'], a['session'], exclude_faculty=faculty_name)
            eligible_map[str(idx)] = eligible

    return render_template('faculty.html',
        faculty_name  = faculty_name,
        my_allocs     = my_allocs,
        total         = len(my_allocs),
        my_swaps      = my_swaps,
        peer_sent     = peer_sent,
        peer_incoming = peer_incoming,
        eligible_map  = eligible_map)

@app.route('/faculty/swap_request', methods=['POST'])
def swap_request():
    if session.get('role') != 'faculty':
        return redirect(url_for('login'))
    faculty_name = session.get('user')
    alloc_index  = int(request.form.get('alloc_index'))
    reason       = request.form.get('reason', '').strip()

    allocs = STATE.get('allocations') or []
    if alloc_index >= len(allocs) or allocs[alloc_index]['faculty'] != faculty_name:
        flash('Invalid allocation selected.', 'error')
        return redirect(url_for('faculty_dashboard'))

    existing = next((s for s in STATE['swap_requests']
                     if s['alloc_index'] == alloc_index and s['status'] == 'pending'), None)
    if existing:
        flash('A swap request for this duty already exists and is pending.', 'error')
        return redirect(url_for('faculty_dashboard'))

    alloc = allocs[alloc_index]
    STATE['swap_requests'].append({
        'id':          str(uuid.uuid4())[:8],
        'requester':   faculty_name,
        'alloc_index': alloc_index,
        'date':        alloc['date'],
        'day':         alloc['day'],
        'session':     alloc['session'],
        'reason':      reason or 'No reason provided',
        'status':      'pending',
        'replaced_by': None,
        'submitted':   datetime.now().strftime('%d-%m-%Y %H:%M'),
        'type':        'admin',
    })
    flash(f'Swap request submitted for {alloc["date"]} {alloc["session"]} session.', 'success')
    return redirect(url_for('faculty_dashboard'))

if __name__ == '__main__':
    init_db()
    app.run(debug=True, port=5050)